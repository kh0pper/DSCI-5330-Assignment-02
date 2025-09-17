#!/usr/bin/env python3
"""
Create Works Cited document for Ford Motor Company Financial Analysis
Version 2: With hyperlinked text instead of showing URLs
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import qn
import os
from datetime import datetime

def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph in Word document
    """
    # This approach creates a hyperlink within the docx structure
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create new run with hyperlink formatting
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add blue color and underline for hyperlink style
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Add font settings
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')  # 12pt * 2
    rPr.append(sz)

    new_run.append(rPr)
    new_run.text = text

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def create_works_cited():
    """Create the works cited document"""
    doc = Document()

    # Set default document formatting
    for style in doc.styles:
        if hasattr(style, 'font'):
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)

    # Set margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Create header with Ford logo
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.exists('Ford_Motor_Company_Logo.png'):
        header_run = header_para.add_run()
        header_run.add_picture('Ford_Motor_Company_Logo.png', width=Inches(2))

    # Add title
    title = doc.add_heading('WORKS CITED', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # Ford Blue

    # Add subtitle
    subtitle = doc.add_paragraph('Ford Motor Company Financial Analysis (2015-2024)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(64, 64, 64)
    subtitle_run.font.italic = True

    doc.add_paragraph()

    # Base GitHub URL for raw files
    base_url = "https://github.com/kh0pper/DSCI-5330-Assignment-02/blob/main"

    # Annual Reports section
    annual_heading = doc.add_heading('Annual Reports', level=1)
    annual_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    annual_reports = [
        ("Ford Motor Company. (2016). 2016 Annual Report.",
         f"{base_url}/Annual%20Report/2016-annual-report.pdf"),
        ("Ford Motor Company. (2017). 2017 Annual Report.",
         f"{base_url}/Annual%20Report/Final-Annual-Report-2017.pdf"),
        ("Ford Motor Company. (2018). 2018 Annual Report.",
         f"{base_url}/Annual%20Report/2018-Annual-Report.pdf"),
        ("Ford Motor Company. (2019). 2019 Annual Report.",
         f"{base_url}/Annual%20Report/Ford-2019-Printed-Annual-Report.pdf"),
        ("Ford Motor Company. (2020). 2020 Annual Report.",
         f"{base_url}/Annual%20Report/Ford-2020-Annual-Report-April-2020.pdf"),
        ("Ford Motor Company. (2021). 2021 Annual Report.",
         f"{base_url}/Annual%20Report/Ford-2021-Annual-Report.pdf"),
        ("Ford Motor Company. (2022). 2022 Annual Report.",
         f"{base_url}/Annual%20Report/2022-Annual-Report-1.pdf"),
        ("Ford Motor Company. (2023). 2023 Annual Report.",
         f"{base_url}/Annual%20Report/2023-Ford-Annual-Report.pdf"),
        ("Ford Motor Company. (2024). 2024 Annual Report.",
         f"{base_url}/Annual%20Report/Ford-2024-Annual-Report.pdf"),
    ]

    for citation, url in annual_reports:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # Add citation text
        citation_run = para.add_run(citation + " ")
        citation_run.font.size = Pt(12)
        citation_run.font.name = 'Times New Roman'

        # Add hyperlinked text instead of URL
        add_hyperlink(para, "[View Report]", url)

    doc.add_paragraph()

    # 10-K Reports section
    ten_k_heading = doc.add_heading('SEC Form 10-K Annual Reports', level=1)
    ten_k_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    ten_k_reports = [
        ("Ford Motor Company. (2016, February 11). Form 10-K for fiscal year ended December 31, 2015. U.S. Securities and Exchange Commission.",
         f"{base_url}/10k/2016_10K_for%20Year%20End%202015%20-%20filed%2002.11.16.pdf"),
        ("Ford Motor Company. (2017, February 9). Form 10-K for fiscal year ended December 31, 2016. U.S. Securities and Exchange Commission.",
         f"{base_url}/10k/2017_10K_for%20Year%20End%202016%20-%20filed%2002.09.17.pdf"),
        ("Ford Motor Company. (2018, February 8). Form 10-K for fiscal year ended December 31, 2017. U.S. Securities and Exchange Commission.",
         f"{base_url}/10k/2018_10K_for%20Year%20End%202017%20-%20filed%2002.08.18.pdf"),
        ("Ford Motor Company. (2019, February 21). Form 10-K for fiscal year ended December 31, 2018. U.S. Securities and Exchange Commission.",
         f"{base_url}/10k/2019_10K_for%20Year%20End%202018%20-%20filed%2002.21.19.pdf"),
        ("Ford Motor Company. (2020, February 5). Form 10-K for fiscal year ended December 31, 2019. U.S. Securities and Exchange Commission.",
         f"{base_url}/10k/2020_10K_for%20Year%20End%202019%20-%20filed%2002.05.20.pdf"),
        ("Ford Motor Company. (2021, February 5). Form 10-K for fiscal year ended December 31, 2020. U.S. Securities and Exchange Commission.",
         f"{base_url}/10k/2021_10K_for%20Year%20End%202020%20-%20filed%2002.05.21.pdf"),
        ("Ford Motor Company. (2022, February 4). Form 10-K for fiscal year ended December 31, 2021. U.S. Securities and Exchange Commission.",
         f"{base_url}/10k/2022_10K_for%20Year%20End%202021%20-%20filed%2002.04.22.pdf"),
        ("Ford Motor Company. (2023, February 3). Form 10-K for fiscal year ended December 31, 2022. U.S. Securities and Exchange Commission.",
         f"{base_url}/10k/2023_10K_for%20Year%20End%202022%20-%20filed%2002.03.23.pdf"),
        ("Ford Motor Company. (2024, February 7). Form 10-K for fiscal year ended December 31, 2023. U.S. Securities and Exchange Commission.",
         f"{base_url}/10k/2024_10K_for%20Year%20End%202023%20-%20filed%2002.07.24.pdf"),
        ("Ford Motor Company. (2025, February 6). Form 10-K for fiscal year ended December 31, 2024. U.S. Securities and Exchange Commission.",
         f"{base_url}/10k/2025_10K_for%20Year%20End%202024%20-%20filed%2002.06.25.pdf"),
    ]

    for citation, url in ten_k_reports:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # Add citation text
        citation_run = para.add_run(citation + " ")
        citation_run.font.size = Pt(12)
        citation_run.font.name = 'Times New Roman'

        # Add hyperlinked text instead of URL
        add_hyperlink(para, "[View 10-K]", url)

    doc.add_paragraph()

    # Financial Data Sources section
    data_heading = doc.add_heading('Financial Data Sources', level=1)
    data_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    data_sources = [
        ("Ford Motor Company. (2024). Ford 10-K Financial Ratios 2015-2024 [Excel spreadsheet]. Internal analysis compilation.",
         f"{base_url}/Ford_10K_Financial_Ratios_2015_2024.xlsx"),
    ]

    for citation, url in data_sources:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # Add citation text
        citation_run = para.add_run(citation + " ")
        citation_run.font.size = Pt(12)
        citation_run.font.name = 'Times New Roman'

        # Add hyperlinked text instead of URL
        add_hyperlink(para, "[View Data]", url)

    doc.add_paragraph()

    # Note section
    note_heading = doc.add_heading('Note', level=1)
    note_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    note_para = doc.add_paragraph(
        "All financial reports and data sources are publicly available documents filed with the "
        "U.S. Securities and Exchange Commission or published by Ford Motor Company. Links provided "
        "direct to the archived versions maintained in the project repository for research consistency "
        "and accessibility."
    )
    note_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    note_para.runs[0].font.size = Pt(12)
    note_para.runs[0].font.italic = True

    doc.add_paragraph()

    # Footer
    footer_para = doc.add_paragraph(
        f'Prepared by: Financial Analysis Team | {datetime.now().strftime("%B %d, %Y")}'
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(10)
    footer_para.runs[0].font.color.rgb = RGBColor(64, 64, 64)

    # Apply document-wide formatting
    for paragraph in doc.paragraphs:
        if not paragraph.style.name.startswith('Heading'):
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            if len(paragraph.text) > 0:
                for run in paragraph.runs:
                    if run.font.size is None or run.font.size == Pt(11):
                        run.font.size = Pt(12)
                        run.font.name = 'Times New Roman'

    # Save the document
    doc.save('work-cited.docx')
    print("Works cited document created successfully with hyperlinked text: work-cited.docx")

if __name__ == "__main__":
    create_works_cited()